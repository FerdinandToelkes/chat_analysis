import os 
import docx
import re
import logging
import nltk
import yaml

from empath import Empath

from chat_analysis.utils import get_sentiment_key_words, get_chat_and_keywords_file_paths


# Configure logging at the module level
logging.basicConfig(
    level=logging.DEBUG, 
    format="%(name)s - %(asctime)s - %(levelname)s - %(message)s"
)

# Get the logger for this script (one logger per module)
logger = logging.getLogger(__name__)





class ChatAnalyzer:
    """ Read in .docx file and analyze chat data """
    def __init__(self, chat_file_path: str, keywords_file_path: str):
        self.file = docx.Document(chat_file_path)
        self.name_one, self.name_two = self.extract_names()
        self.chat = self.extract_conversation_parts()
        self.empath = Empath()
        self.sentiment_keywords = get_sentiment_key_words(keywords_file_path)
        

    def print_chat_file(self):
        for para in self.file.paragraphs:
            print(para.text)
            print("")

    def print_responses(self, name: str):
        print(f"Responses from {name}:")
        print("")
        for response in self.chat[name]:
            print(response)
            print("")

    def extract_names(self) -> tuple[str, str]:
        """ Extract names of the two people in the chat which are in a line with the pattern {name}{month}/{day}/{year} 
        with name being a string to be extracted. It is assumed that the names are at the beginning of the paragraphs. 
        Example: 
        I think that's exactly what's happening to you. It's like you're describing a sense of disconnection ...
        Alex01/22/2025
        be surrounded by people and still feel alone?
        you01/22/2025
        -> Alex, you
        """ 
        # ^ asserts the start of a line, [A-Za-z]+ matches one or more letters, \d{n} matches n digits, / matches a literal slash
        name_pattern = re.compile(r"^([A-Za-z]+)\d{2}/\d{2}/\d{4}$")
        name_one, name_two = "", ""
        logger.debug("Extracting names ...")

        for para in self.file.paragraphs:
            
            match = name_pattern.match(para.text)
            if match:
                # Extract name e.g. from Alex01/22/2025
                name = match.group(1) # group is defined by the parantheses in the regex pattern
                if not name_one:
                    name_one = name
                    logger.debug(f"Name one is set to '{name_one}'")
                elif name != name_one and not name_two:
                    name_two = name
                    logger.debug(f"Name two is set to '{name_two}'")
                elif name_two != "" and name_one != "":
                    break
        
        if not name_one or not name_two:
            raise ValueError("Could not extract two valid names from the document.")
            
        return name_one, name_two

    

    def extract_conversation_parts(self) -> dict:
        """ Extract conversation parts from chat file assuming that they are 
        separated by the pattern '{name}{month}/{day}/{year}', where name 
        is either name_one or name_two.

        Returns:
            chat (dict): dictionary with keys name_one and name_two and values being lists of strings
        ----------------
        Example:
        I think that's exactly what's happening to you. It's like you're describing a sense of disconnection ...
        Alex01/22/2025
        be surrounded by people and still feel alone?
        you01/22/2025
        -> chat = {"Alex": ["I think that's exactly what's happening to you. It's like you're describing a sense of disconnection ...", "be surrounded by people and still feel alone?"], "you": ["be surrounded by people and still feel alone?"]}
        """
        # Initialize chat dictionary and unwanted texts
        chat = {self.name_one: [], self.name_two: []}
        unwanted_text = ["Replika is an AI and cannot provide medical advice. In a crisis, seek expert help."]

        # Regex pattern to match a line containing a name and a date
        # ^ asserts the start of a line, [A-Za-z]+ matches one or more letters, \d{n} matches n digits, / matches a literal slash
        name_pattern = re.compile(r"^([A-Za-z]+)\d{2}/\d{2}/\d{4}$")
        
        # Track last speaker assuming the last message is from name_one
        previous_name = self.name_one
        name_changed = False

        logger.debug("Extracting conversation parts ...")
        for para in self._read_paragraphs_reversed():
            # Skip all unwanted text
            if para in unwanted_text:
                continue

            # either name date line or message
            match = name_pattern.match(para) 
            if match:
                # Extract name e.g. from Alex01/22/2025
                name = match.group(1) # group is defined by the parantheses in the regex pattern
                if name not in chat:
                    raise ValueError(f"Name {name} does not match any of the names in the chat")
                
                # boolean to track speaker change
                name_changed = name != previous_name 
                previous_name = name
                continue # Skip name date line
            
            # Append or merge based on whether the speaker changed
            chat = self._add_text_to_chat(chat, previous_name, para, name_changed)

        # Reverse the order of the messages
        for person in chat:
            chat[person] = chat[person][::-1]

        return chat

    def _read_paragraphs_reversed(self):
        """ Read paragraphs in reverse order """
        # Filter out empty paragraphs
        paragraphs = [para.text for para in self.file.paragraphs if para.text.strip()]
        return reversed(paragraphs)  # Returns a generator

    def _add_text_to_chat(self, chat: dict, name: str, para: str, name_changed: bool) -> dict:
        """ Append or merge text to chat based on whether the speaker changed 
        
        Args:
            chat (dict): dictionary containing chat messages
            name (str): name of the speaker
            para (str): text to add to chat
            name_changed (bool): boolean indicating whether the speaker has changed

        Returns:
            chat (dict): updated chat dictionary
        """
        if name_changed or not chat[name]:
            chat[name].append(para)
        else:
            # note that we are iterating through the text in reversed order
            merged_text = para + " " + chat[name][-1]
            chat[name][-1] = merged_text
        return chat
        

    def get_basic_chat_info(self) -> dict:
        """ Get basic info of chat such as number of messages and number of words for each person in the chat and in total.
        
        Returns:
            stats (dict): dictionary containing basic chat statistics
        """
        stats = {}
        stats["name_one"] = self.name_one
        stats["name_two"] = self.name_two
        stats["number_of_responses_name_one"] = len(self.chat[self.name_one])
        stats["number_of_responses_name_two"] = len(self.chat[self.name_two])
        stats["total_number_of_seperate_texts"] = stats["number_of_responses_name_one"] + stats["number_of_responses_name_two"]
        stats["total_number_of_words_name_one"] = sum([self._count_words(text) for text in self.chat[self.name_one]])
        stats["total_number_of_words_name_two"] = sum([self._count_words(text) for text in self.chat[self.name_two]])
        stats["total_number_of_words"] = stats["total_number_of_words_name_one"] + stats["total_number_of_words_name_two"]
        return stats
    
    def _count_words(self, text: str) -> int:
        """ Count the number of words in a text 
        
        Args:
            text (str): text to count words in

        Returns:
            num_words (int): number of words in the text
        """
        words = text.split()
        words = [word for word in words if word != "-"]
        return len(words)
    
    def get_chat_sentiment(self):
        """ Get the sentiment of the chat using predefined sentiment keywords. """
        sentiment = {}
        for person in [self.name_one, self.name_two]:
            messages = " ".join(self.chat[person]) if self.chat[person] else ""
            sentiment[person] = self.empath.analyze(messages, categories=self.sentiment_keywords, normalize=True) if messages else {}
        return sentiment


    def get_chat_wordcloud(self, chat):
        pass

    def get_chat_graph(self, chat):
        pass

    def get_chat_summary(self, chat):
        pass






def main():
    chat_file_path, keywords_file_path = get_chat_and_keywords_file_paths()
    analyzer = ChatAnalyzer(chat_file_path, keywords_file_path)
    # analyzer.print_chat_file()
    print(f"Name one: {analyzer.name_one}")
    print(f"Name two: {analyzer.name_two}")
    print(f"Chat: {analyzer.chat}")
    print(analyzer.get_basic_chat_info())
    #print(analyzer.get_chat_sentiment())
    

if __name__ == '__main__':
    main()
    